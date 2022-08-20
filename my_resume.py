from pydoc import Doc
from docx import Document
from docx.shared import Inches

document = Document()


document.add_heading('CHAIMAA HADDAR', 0)

# location phone number and email
location = 'Houston, Tx'
phone_number = '832-998-0860'
email = 'chaimaahaddar@hotmail.com'

document.add_paragraph(
    location + ' | ' + phone_number + ' | ' + email)

#about me
p = document.add_paragraph()
document.add_heading('ABOUT ME')
about_me =  ('Software QA Tester with full system development experience, including designing, developing, and implementing test plans and executing testing strategies for Web-based and client/server applications using manual and automation testing.')

document.add_paragraph(about_me)
p = document.add_paragraph()

#education
document.add_heading('EDUCATION')
p = document.add_paragraph()

#education1
university1 = ('Houston Community College')
from_date = ('JUNE 2021')
to_date = ('Present')

p.add_run(university1 + '\n').bold = True 
p.add_run(from_date + ' - ' + to_date + '\n').italic = True
p.add_run("Associate's Degree in Computer Science\r") 
p.add_run('\n')

#education2
university2 = ('Harvard University Online')
from_date = ('JUNE 2020')
to_date = ('DECEMBER 2020')

p.add_run(university2 + '\n').bold = True 
p.add_run(from_date + ' - ' + to_date + '\n').italic = True
p.add_run('Certification online course Computer Science: CS50 \r') 
p.add_run('\n')

document.add_paragraph('\r')

#education3
university3 = ('University of Verona UNIVR')
from_date = ('SEPTEMBER 2014 ')
to_date = ('JULY 2017')

p.add_run(university3 + '\n').bold = True
p.add_run(from_date + ' - ' + to_date + '\n').italic = True
p.add_run("Bachelor's Degree in Foreign Languages and Literatures \r")


#work experience 1
document.add_heading('WORK EXPERIENCE')
document.add_heading('Hp')
p = document.add_paragraph()


jobtitle = ('QA Engineer')
from_date = ('JULY 2021')
to_date = ('PRESENT')

p.add_run(jobtitle + '\n').bold = True 
p.add_run(from_date + ' - ' + to_date + '\n').italic = True 
p.add_run('\n')


p.add_run('Strong knowledge of software QA methodologies, tools, and processes. \n')
p.add_run('Finding bugs and collaborating with developers on fixes to various issues before the launch of the product, app or software. \n')
p.add_run('Performing various testing tools, such as Selenium. \n')
p.add_run('Accessing and Inspecting Log Files with the use of Event Viewer and Fiddler. \n')
p.add_run('Performing localization duties on several Hp apps making sure software features are properly localized. \n')
p.add_run('Software Benchmarking to set products performances. \n')



document.add_heading('Koch Industries - Optimized Process Designs')
p = document.add_paragraph()
document.add_paragraph()

#work experience 2

company = ('IT Operations Analyst')
from_date = ('JANUARY 2020')
to_date = ('DECEMBER 2020')

p.add_run(company + '\n').bold = True 
p.add_run(from_date + ' - ' + to_date + '\n').italic = True 
p.add_run('\n')

p.add_run('Handling data in different methodologies: managing the tech inventory and data analysis with the use of advanced software programs such as Sql, Python, Ms Excell and various Data Visualization tools (Power Pivot, Query, Bi).\n')
p.add_run('Offering technician support to employees\n')
p.add_run('Responsible of data backups and system security operations, like user authorizations and firewalls. \n')



#Table of skills

document.add_heading('SKILLS', level=1)
document.add_paragraph('Experience Tools', style='Intense Quote')


records = (
    ('Selenium', 'PyTest', 'Postman'),
    ('Jira', 'Azure DevOps', '.Net')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'MySQL'
hdr_cells[1].text = 'Python'
hdr_cells[2].text = 'JSON'
for MySQL, Python, JSON in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(MySQL)
    row_cells[1].text = Python
    row_cells[2].text = JSON




document.save('resume.docx')
